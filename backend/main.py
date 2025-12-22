import traceback
from fastapi import FastAPI, File, Form, UploadFile, HTTPException, Depends, APIRouter, status, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials, OAuth2PasswordRequestForm
from pydantic import BaseModel, EmailStr
from openai import OpenAI
import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import shutil
import json
import uuid
from pathlib import Path
from passlib.context import CryptContext
from dotenv import load_dotenv
from backend.enhanced_resume_parser import EnhancedResumeParser
from backend.database import Resume, create_tables, get_db, ChatSession, ChatMessage, ResumeAnalysis, User, UserProfile, UserActivity, JobPosting, UserJobPreferences, JobApplication, SavedJob, JobMatch,get_db,Base, User, UserProfile, JobPosting, JobApplication, SavedJob, UserJobPreferences,JobMatch,Resume,JobAlert
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import tempfile
from datetime import datetime
from backend.auth import verify_password, get_password_hash, create_access_token, get_current_user_id
from typing import Optional
from backend.job_api_service import JobAPIService
from backend.job_matching import JobMatchingEngine
from typing import List, Optional
from datetime import datetime, timedelta
from sqlalchemy import desc, and_, or_
from backend.email_service import send_job_alert_email
from apscheduler.schedulers.background import BackgroundScheduler
import atexit



security = HTTPBearer(auto_error=False)
job_api_service = JobAPIService()
matching_engine = JobMatchingEngine()


async def get_current_user_optional(
    authorization: Optional[str] = Header(None),
    db: Session = Depends(get_db)
) -> Optional[int]:
    if not authorization:
        return None
    
    try:
        token = authorization.replace("Bearer ", "")
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id = payload.get("user_id") or payload.get("sub")
        return int(user_id) if user_id else None
    except:
        return None



load_dotenv()

app = FastAPI()

#Creating database table
create_tables()

#Enabling CORS for frontend connection
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://ai-career-mentor-kappa.vercel.app",
        "https://www.careermentorlab.com",
        "https://careermentorlab.com"],
      # Next.js default port
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"]
)

client = OpenAI(api_key= os.getenv("OPENAI_API_KEY"))

# Initializing resume parser
resume_parser = EnhancedResumeParser()

#Creating Uploads directory
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

# Create uploads directory
UPLOAD_DIR = "uploads/resumes"
os.makedirs(UPLOAD_DIR, exist_ok=True)

class ChatMessageSchema(BaseModel):
    message: str
    session_id: str | None = None

class CoverLetterRequest(BaseModel):
    resume_data: dict
    job_description: str
    company_name: str
    position_title: str
    tone: str = "professional" # professional, friendly, enthusiastic

class ResumeOptimizationRequest(BaseModel):
    resume_data : dict
    job_description : str
    company_name : str
    position_title: str
    optimization_level: str = "moderate" # conservative, moderate, aggressive

class DocxGenerationRequest(BaseModel):
    content: str
    filename: str
    doc_type: str
    company_name: str
    position_title: str

class UserCreate(BaseModel):
    username: str
    email: str
    password: str
    full_name: str
    role: str ="job_seeker"

class UserLogin(BaseModel):
    username: str
    password: str


class UserResponse(BaseModel):
    id:int
    username: str
    email: str
    full_name: str
    role: str
    is_active: bool
    created_at: datetime
    resume_analyses_count: int
    cover_letters_count: int
    optimizations_count: int
    chat_messages_count: int

    class Config:
        from_attributes = True

class ProfileUpdate(BaseModel):
    current_role: str = None
    experience_level: str = None
    industry: str = None
    location: str = None
    years_of_experience: Optional[int] = None
    career_goals: Optional[str] = None
    salary_range: str = None
    technical_skills: list = None
    soft_skills: list = None
    certifications: list = None
    languages: list = None
    preferred_roles: list = None
    preferred_companies: list = None
    remote_preference: str = None
    willing_to_relocate: bool = None
    linkedin_url: str = None
    github_url: str = None
    portfolio_url: str = None

class JobPreferencesCreate(BaseModel):
    desired_roles: Optional[List[str]] = []
    preferred_locations: Optional[str] = None  # JSON string
    remote_preference: str = "flexible"  # remote_only, flexible, onsite
    minimum_salary: Optional[int] = None
    maximum_salary: Optional[int] = None
    preferred_companies: Optional[str] = None  # JSON string
    company_sizes: Optional[str] = None  # JSON string
    willing_to_relocate: bool = False

class JobApplicationCreate(BaseModel):
    job_id: int
    cover_letter: Optional[str] = None
    resume_version: Optional[str] = None

class JobApplicationUpdate(BaseModel):
    status: Optional[str] = None
    notes: Optional[str] = None
    interview_date: Optional[str] = None



class ProfileResponse(BaseModel):
    current_role: Optional[str] = None
    industry: Optional[str] = None
    years_of_experience: Optional[int] = None
    career_goals: Optional[str] = None
    location: Optional[str] = None
    updated_at: datetime

    class Config:
        from_attributes = True


#Helper fuctiion for my experience level problem
def extract_experience_level(title: str, description: str = "") -> str:
    """
    Extract experience level from job title and description
    Returns: 'Junior', 'Mid', or 'Senior'
    """
    text = (title + " " + description).lower()
    
    # Senior level keywords
    senior_keywords = [
        'senior', 'sr.', 'lead', 'principal', 'staff', 'architect',
        'head of', 'director', 'manager', 'expert', 'specialist',
        '5+ years', '5-', '6+ years', '7+ years', '8+ years',
        'experienced', 'advanced'
    ]
    
    # Junior level keywords
    junior_keywords = [
        'junior', 'jr.', 'entry', 'graduate', 'intern',
        'associate', 'trainee', '0-2 years', '1-2 years',
        'early career', 'beginner', 'starter'
    ]
    
    # Check for senior
    for keyword in senior_keywords:
        if keyword in text:
            return 'Senior'
    
    # Check for junior
    for keyword in junior_keywords:
        if keyword in text:
            return 'Junior'
    
    # Default to Mid if no specific level found
    return 'Mid'

#Admin route
@app.get("/api/admin/check-config")
async def check_api_configuration():
    """Check if API keys are configured"""
    return {
        "adzuna_configured": bool(job_api_service.adzuna_app_id and job_api_service.adzuna_api_key),
        "jsearch_configured": bool(job_api_service.jsearch_api_key),
        "adzuna_app_id": job_api_service.adzuna_app_id[:10] + "..." if job_api_service.adzuna_app_id else None,
        "adzuna_key_length": len(job_api_service.adzuna_api_key) if job_api_service.adzuna_api_key else 0,
        "jsearch_key_length": len(job_api_service.jsearch_api_key) if job_api_service.jsearch_api_key else 0
    }

@app.post("/api/admin/populate-jobs")
async def populate_initial_jobs(db: Session = Depends(get_db)):
    """Populate database with initial jobs - NO AUTH REQUIRED FOR TESTING"""
    
    queries = [
        "software developer",
        "frontend developer", 
        "backend developer",
        "full stack developer",
        "python developer",
        "junior developer",
        "senior developer",
        "lead developer",
         # I would add more soont:
    ]
    
    total_added = 0
    total_updated = 0
    
    for query in queries:
        print(f"Fetching jobs for: {query}")
        jobs_data = job_api_service.fetch_and_store_jobs(query, "United Kingdom", 20)
        
        for job_data in jobs_data:
            existing = db.query(JobPosting).filter(
                JobPosting.external_id == job_data['external_id']
            ).first()
            
            if not existing:
                skills = job_api_service.extract_skills_from_description(
                    job_data.get('description', '')
                )

                #Extracting Experience level
                experience_level = extract_experience_level(
                    job_data['title'],
                    job_data.get('description', '')
                )
                
                job = JobPosting(
                    title=job_data['title'],
                    company_name=job_data['company_name'],
                    company_logo_url=job_data.get('company_logo_url'),
                    location=job_data['location'],
                    remote_type=job_data['remote_type'],
                    description=job_data['description'],
                    requirements=job_data.get('requirements', ''),
                    salary_min=job_data.get('salary_min'),
                    salary_max=job_data.get('salary_max'),
                    salary_currency='GBP',
                    experience_level=experience_level,
                    employment_type=job_data['employment_type'],
                    required_skills=skills,
                    external_id=job_data['external_id'],
                    source=job_data['source'],
                    apply_url=job_data.get('apply_url'),
                    posted_date=job_data.get('posted_date') or datetime.utcnow(),
                    is_active=True
                )
                db.add(job)
                total_added += 1
            else:
                experience_level = extract_experience_level(
                    job_data['title'],
                    job_data.get('description', '')
                )
                existing.experience_level = experience_level
                total_updated += 1
        
        db.commit()
        print(f"  Added: {total_added}, Updated: {total_updated}")
    
    return {
        "status": "success",
        "jobs_added": total_added,
        "jobs_updated": total_updated,
        "total": total_added + total_updated
    }    

@app.get("/")
async def root():
    return {"message": "AI Career Mentor API is running"}

@app.post("/api/chat")
async def chat_with_ai(
    chat_message: ChatMessageSchema, 
    db: Session = Depends(get_db),
    current_user_id: Optional[int] = Depends(get_current_user_optional)
):
    try:
        # Creating or getting session
        session_id = chat_message.session_id or str(uuid.uuid4())

        # Getting or creating session
        chat_session = db.query(ChatSession).filter(ChatSession.session_id == session_id).first()
        if not chat_session:
            chat_session = ChatSession(
                session_id=session_id,
                user_id=current_user_id  # This will be None if not authenticated
            )
            db.add(chat_session)
            db.commit()
            db.refresh(chat_session)

        # Save user message
        user_message = ChatMessage(
            session_id=session_id,
            role="user",
            content=chat_message.message
        )
        db.add(user_message)
        db.commit()
        db.refresh(user_message)

        # Get conversation history (last 10 messages)
        recent_messages = db.query(ChatMessage).filter(
            ChatMessage.session_id == session_id
        ).order_by(ChatMessage.timestamp.desc()).limit(10).all()

        # Build conversation context
        conversation_history = []
        for msg in reversed(recent_messages):
            role = "assistant" if msg.role == "ai" else msg.role
            conversation_history.append({
                "role": role,
                "content": msg.content
            })

        # System prompt
        system_prompt = """You are an AI career mentor specializing in tech careers.
        Help with resume advice, career transitions, skill development, interview prep,
        and job search strategies in the tech industry. Be practical and actionable.
        Remember previous context in this conversation."""

        # Prepare messages for OpenAI
        messages = [{"role": "system", "content": system_prompt}] + conversation_history

        # Call OpenAI
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=messages,
            max_tokens=500,
            temperature=0.7
        )

        ai_response = response.choices[0].message.content

        # Save AI response
        ai_message = ChatMessage(
            session_id=session_id,
            role="ai",
            content=ai_response
        )
        db.add(ai_message)

        # Update session title if it's the first message
        if chat_session.title == "New Conversation":
            title = (
                chat_message.message[:50] + "..."
                if len(chat_message.message) > 50
                else chat_message.message
            )
            chat_session.title = title

        # Increment user's chat message count if authenticated
        if current_user_id:
            user = db.query(User).filter(User.id == current_user_id).first()
            if user:
                user.chat_messages_count = (user.chat_messages_count or 0) + 1

        db.commit()

        return {
            "response": ai_response,
            "session_id": session_id,
            "status": "success"
        }

    except Exception as e:
        db.rollback()
        return {
            "response": f"Sorry, there was an error: {str(e)}",
            "status": "error"
        }
    
@app.get("/api/chat/sessions")
async def get_chat_sessions(
    db: Session = Depends(get_db),
    current_user_id: Optional[int] = Depends(get_current_user_optional)
):
    """Get chat sessions for current user only"""
    if current_user_id:
        # Get sessions for authenticated user
        sessions = db.query(ChatSession).filter(
            ChatSession.user_id == current_user_id
        ).order_by(ChatSession.updated_at.desc()).all()
    else:
        # For non-authenticated users, return empty list or sessions without user_id
        sessions = db.query(ChatSession).filter(
            ChatSession.user_id.is_(None)
        ).order_by(ChatSession.updated_at.desc()).all()
    
    return [
        {
            "session_id": session.session_id,
            "title": session.title,
            "created_at": session.created_at.isoformat(),
            "updated_at": session.updated_at.isoformat()
        }
        for session in sessions
    ]


@app.get("/api/chat/sessions/{session_id}")
async def get_chat_history(session_id: str, db: Session = Depends(get_db)):
    """Get chat history for a specific session"""
    messages = db.query(ChatMessage).filter(
        ChatMessage.session_id == session_id
    ).order_by(ChatMessage.timestamp.asc()).all()
    
    return [
        {
            "role": msg.role,
            "content": msg.content,
            "timestamp": msg.timestamp.isoformat()
        }
        for msg in messages
    ]

@app.delete("/api/chat/sessions/{session_id}")
async def delete_chat_session(session_id: str, db: Session = Depends(get_db)):
    """Delete a chat session and all its messages"""
    chat_session = db.query(ChatSession).filter(ChatSession.session_id == session_id).first()
    if chat_session:
        db.delete(chat_session)
        db.commit()
        return {"status": "success", "message": "Session deleted"}
    return {"status": "error", "message": "Session not found"} 
   
@app.post("/api/upload-resume")
async def upload_resume(file: UploadFile = File(...)):
    """Uploading and parsing resume file"""
    try:
        # Validating file type
        if not file.filename.lower().endswith(('.pdf', '.docx')):
            raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")
        
        # Create uploads directory if it doesn't exist
        upload_dir = Path("uploads")
        upload_dir.mkdir(exist_ok=True)
        
        # Saving uploaded file
        file_path = upload_dir / file.filename
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Parsing resume
        file_type = file.filename.split('.')[-1].lower()
        parsed_data = resume_parser.parse_resume(str(file_path), file_type)
        
        # Cleaning up uploaded file
        os.remove(file_path)
        
        if "error" in parsed_data:
            raise HTTPException(status_code=500, detail=parsed_data["error"])
        
        return {
            "status": "success",
            "filename": file.filename,
            "data": parsed_data
        }
    except Exception as e:
        # Cleaning up file if it exists
        try:
            if 'file_path' in locals() and Path(file_path).exists():
                os.remove(file_path)
        except:
            pass
        raise HTTPException(status_code=500, detail=f"Error processing resume: {str(e)}")
    
@app.post("/api/analyze-resume")
async def analyze_resume(
    resume_data: dict, 
    db: Session = Depends(get_db),
    current_user_id: Optional[int] = Depends(get_current_user_optional)
    ):
    """Analyzing parsed resume with AI and providing feedback"""
    try:
        # Try to get current user ID if authenticated
        user_id = None
        try:
            user_id = current_user_id
        except:
            pass  # Continue without authentication

        # Create analysis prompt (your existing code)
        skills_text = ""
        for category, skills in resume_data.get("skills", {}).items():
            if skills:
                skills_text += f"{category.title()}: {', '.join(skills)}\n"
        
        analysis_prompt = f"""
        As a tech career expert and ATS specialist, analyze this comprehensive resume data:

        ## Resume Overview
        - **Experience**: {resume_data.get('estimated_experience', 0)} years
        - **Resume Length**: {resume_data.get('word_count', 0)} words
        - **ATS Score**: {resume_data.get('resume_score', {}).get('score', 0)}/100

        ## Contact Information
        - Email: {resume_data.get('contact_info', {}).get('email', 'Missing')}
        - LinkedIn: {resume_data.get('contact_info', {}).get('linkedin', 'Missing')}
        - GitHub: {resume_data.get('contact_info', {}).get('github', 'Missing')}
        - Portfolio: {resume_data.get('contact_info', {}).get('portfolio', 'Missing')}

        ## Technical Skills Found
        {chr(10).join([f"**{category.replace('_', ' ').title()}**: {', '.join(skills)}" 
                    for category, skills in resume_data.get('skills', {}).items() if skills])}

        ## Education
        {chr(10).join([f"- {edu.get('degree', 'N/A')} at {edu.get('institution', 'N/A')} ({edu.get('year', 'N/A')})" 
                    for edu in resume_data.get('education', [])])}

        ## Work Experience
        {chr(10).join([f"- {exp.get('title', 'N/A')} ({exp.get('dates', 'N/A')})" 
                    for exp in resume_data.get('experience', [])])}

        ## Certifications
        {', '.join(resume_data.get('certifications', [])) or 'None found'}

        Please provide a comprehensive analysis with:

        ### 1. 🎯 Overall Assessment
        Rate the resume's strength for tech roles and current market competitiveness.

        ### 2. 📊 ATS Optimization
        Specific improvements for Applicant Tracking Systems based on the {resume_data.get('resume_score', {}).get('score', 0)}/100 score.

        ### 3. 🔧 Technical Skills Gap Analysis
        What skills are missing for 2024/2025 tech market? Which skills need more prominence?

        ### 4. 📝 Content & Structure Improvements
        How to improve work experience descriptions, education section, and overall format.

        ### 5. 🔗 Professional Presence
        Recommendations for LinkedIn, GitHub, and portfolio optimization.

        ### 6. 🚀 Top 3 Priority Actions
        Most important changes to make immediately for better results.

        Keep feedback specific, actionable, and tailored to current tech industry standards.
        """

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an expert tech recruiter and career coach specializing in resume optimization for software engineers, developers, and tech professionals."},
                {"role": "user", "content": analysis_prompt}
            ],
            max_tokens=800,
            temperature=0.7
        )

        # Increment user's resume analysis count if authenticated
        if user_id:
            try:
                user = db.query(User).filter(User.id == int(user_id)).first()
                if user:
                    user.resume_analyses_count = (user.resume_analyses_count or 0) + 1
                    db.commit()
                    print(f"Updated resume count for user {current_user_id}: {user.resume_analyses_count}")
            except Exception as e:
                print(f"Error updating resume analysis count: {e}")
        
        return {
            "status": "success",
            "analysis": response.choices[0].message.content,
            "resume_data": resume_data
        }
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error analyzing resume: {str(e)}")

@app.post("/api/generate-cover-letter")
async def generate_cover_letter(
    request: CoverLetterRequest,
    db: Session = Depends(get_db),
    current_user_id: Optional[int] = Depends(get_current_user_optional)
    ):
    """Generate personalized cover letter based on resume and job description"""
    try:
        # Extract key information from resume
        contact_info = request.resume_data.get('contact_info', {})
        skills = request.resume_data.get('skills', {})
        experience = request.resume_data.get('experience', [])
        education = request.resume_data.get('education', [])
        
        # Build skills summary
        all_skills = []
        for category, skill_list in skills.items():
            all_skills.extend(skill_list)
        skills_text = ', '.join(all_skills[:10])  # Top 10 skills
        
        # Build experience summary
        recent_experience = experience[:2] if experience else []
        experience_text = ""
        for exp in recent_experience:
            experience_text += f"- {exp.get('title', 'N/A')} at {exp.get('company', 'Previous Company')} ({exp.get('dates', 'Recent')})\n"
        
        # Education summary
        education_text = ""
        if education:
            latest_edu = education[0]
            education_text = f"{latest_edu.get('degree', 'Degree')} from {latest_edu.get('institution', 'University')}"
        
        # Set tone-based instructions
        tone_instructions = {
            "professional": "Write in a formal, professional tone. Be concise and focus on qualifications.",
            "friendly": "Write in a warm, approachable tone while maintaining professionalism.",
            "enthusiastic": "Write with enthusiasm and energy, showing genuine excitement for the role."
        }
        
        tone_instruction = tone_instructions.get(request.tone, tone_instructions["professional"])
        
        # Create comprehensive prompt
        cover_letter_prompt = f"""
        Write a personalized, compelling cover letter for a tech professional applying to {request.company_name} for the {request.position_title} position.
        
        {tone_instruction}
        
        ## Candidate Information:
        **Contact**: {contact_info.get('email', 'candidate@email.com')}
        **LinkedIn**: {contact_info.get('linkedin', 'Available upon request')}
        **GitHub**: {contact_info.get('github', 'Available upon request')}
        
        **Technical Skills**: {skills_text}
        
        **Recent Experience**:
        {experience_text or "Recent experience in software development"}
        
        **Education**: {education_text or "Computer Science background"}
        
        ## Job Description:
        {request.job_description}
        
        ## Cover Letter Requirements:
        1. **Professional header** with contact information
        2. **Engaging opening** that mentions the specific role and company
        3. **Skills alignment** - match candidate's skills to job requirements
        4. **Experience highlights** - specific examples relevant to the role
        5. **Company knowledge** - show research about the company (if possible from job description)
        6. **Strong closing** with clear next steps
        
        ## Important Guidelines:
        - Keep to 3-4 paragraphs maximum
        - Use specific examples from the candidate's background
        - Match keywords from the job description for ATS optimization
        - Show genuine interest in the company and role
        - End with a professional call-to-action
        - Do not use placeholder text - make it feel personal and specific
        
        Format as a complete cover letter ready to send.
        """
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system", 
                    "content": "You are an expert career coach and professional writer specializing in creating compelling cover letters for tech professionals. You write personalized, ATS-optimized cover letters that get results."
                },
                {"role": "user", "content": cover_letter_prompt}
            ],
            max_tokens=1000,
            temperature=0.7
        )
        
        cover_letter = response.choices[0].message.content
        
        # Generate additional tips
        tips_prompt = f"""
        Based on this cover letter for a {request.position_title} role at {request.company_name}, provide 3-5 specific tips to make it even stronger:
        
        {cover_letter}
        
        Focus on:
        - ATS optimization opportunities
        - Ways to make it more compelling
        - Industry-specific improvements
        - Quantification opportunities
        """
        
        tips_response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": "You are a cover letter optimization expert. Provide specific, actionable improvement tips."
                },
                {"role": "user", "content": tips_prompt}
            ],
            max_tokens=400,
            temperature=0.6
        )
        
        optimization_tips = tips_response.choices[0].message.content

        # Increment user's cover letter count if authenticated
        if current_user_id:
            try:
                user = db.query(User).filter(User.id == int(current_user_id)).first()
                if user:
                    user.cover_letters_count = (user.cover_letters_count or 0) + 1
                    db.commit()
                    print(f"Updated cover letter count for user {current_user_id}: {user.cover_letters_count}")
            except Exception as e:
                print(f"Error updating cover letter count: {e}")
        
        return {
            "status": "success",
            "cover_letter": cover_letter,
            "optimization_tips": optimization_tips,
            "word_count": len(cover_letter.split()),
            "company_name": request.company_name,
            "position_title": request.position_title,
            "tone": request.tone
        }
    
    except Exception as e:
        return {
            "status": "error",
            "message": f"Error generating cover letter: {str(e)}"
        }
    
    
@app.post("/api/optimize-resume")
async def optimize_resume(
    request: ResumeOptimizationRequest,
    db: Session = Depends(get_db),
    current_user_id: Optional[int] = Depends(get_current_user_optional)
):
    """Optimize resume using professional two-column template structure"""
    try:
        # Extract current resume information
        contact_info = request.resume_data.get('contact_info', {})
        skills = request.resume_data.get('skills', {})
        experience = request.resume_data.get('experience', [])
        education = request.resume_data.get('education', [])
        raw_text = request.resume_data.get('raw_text', '')
        
        # Build current skills summary
        all_skills = []
        for category, skill_list in skills.items():
            all_skills.extend(skill_list)
        current_skills = ', '.join(all_skills)

        # Build experience summary from the candidate's actual experience
        experience_summary = ""
        for exp in experience:
            title = exp.get('title', 'Position')
            company = exp.get('company', 'Company')
            dates = exp.get('dates', 'Recent')
            description = exp.get('description', '')
            experience_summary += f"{title} at {company} ({dates}). {description[:100]}... "

        # If no experience found, use a default
        if not experience_summary:
            experience_summary = "Software development experience with various technologies"

        
        # Create professional template optimization prompt
        optimization_prompt = f"""
        You are an expert resume writer. Create a concise, 1-page optimized resume for {request.position_title} at {request.company_name} using the ACTUAL candidate information provided.

        ## CANDIDATE'S ACTUAL INFORMATION:
        Name: Extract from resume or use "Professional Candidate" if not clear
        Email: {contact_info.get('email', 'email@example.com')}
        Phone: {contact_info.get('phone', 'phone number')}
        LinkedIn: {contact_info.get('linkedin', 'Available on request')}
        GitHub: {contact_info.get('github', 'Available on request')}

        Current Skills: {current_skills}
        Experience: {experience_summary[:500] if experience_summary else 'Software development experience'}
        Education: {', '.join([f"{edu.get('degree', 'Degree')} from {edu.get('institution', 'University')}" for edu in education]) if education else 'Computer Science education'}

        ## TARGET JOB REQUIREMENTS:
        {request.job_description[:1000]}

        ## OUTPUT REQUIREMENTS:
        - Use REAL candidate information, not placeholders
        - Maximum 600 words total
        - Professional Summary: 2-3 lines targeting this specific role
        - Experience: Use candidate's ACTUAL job titles and companies
        - Projects: Use candidate's ACTUAL projects
        - Skills: Use candidate's ACTUAL skills, prioritized for this job

        ## TEMPLATE:

        **[Use actual candidate name or "PROFESSIONAL CANDIDATE"]**

        **{contact_info.get('email', 'email@example.com')} | {contact_info.get('phone', 'phone')} | United Kingdom | LinkedIn | GitHub**

        **PROFESSIONAL SUMMARY**
        [2-3 lines specifically targeting {request.position_title} using candidate's actual background and skills relevant to this job]

        **CORE SKILLS**
        [Organize candidate's actual skills by relevance to job posting]

        **EXPERIENCE** 
        [Use candidate's ACTUAL job history - select 2-3 most relevant positions with real titles, companies, and dates]

        **PROJECTS**
        [Use candidate's ACTUAL projects - select 2 most relevant with real project names and technologies]

        **EDUCATION**
        [Use candidate's actual education with real degree, institution, and relevant modules]

        Use the candidate's REAL information throughout - no placeholders or generic content.
        """
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system", 
                    "content": "You are an expert resume writer specializing in creating professional, ATS-optimized resumes using structured templates. You excel at matching candidate backgrounds to specific job requirements while maintaining formatting consistency and professional presentation."
                },
                {"role": "user", "content": optimization_prompt}
            ],
            max_tokens=2500,  # Increased for comprehensive template
            temperature=0.5   # Lower temperature for more consistent formatting
        )

        optimized_resume = response.choices[0].message.content
        
        # Generate detailed optimization summary
        changes_prompt = f"""
        Analyze the resume optimization for {request.position_title} at {request.company_name}:

        ## Original Resume Summary:
        - Skills: {current_skills[:200]}...
        - Experience: {len(experience)} positions
        - Education: {len(education)} qualifications

        ## Optimization Results:
        Provide a summary of:

        **1. Template Improvements**
        - Applied professional two-column layout
        - Enhanced visual hierarchy and readability
        - Optimized section organization for ATS scanning

        **2. Content Optimization**
        - Keywords added from job description
        - Skills reordered by relevance to target role
        - Quantified achievements where possible
        - Enhanced professional summary for target position

        **3. ATS Enhancements**
        - Improved keyword density for target role
        - Standardized formatting for ATS compatibility
        - Clear section headers and structure

        **4. Match Score Improvement**
        - Estimated increase in job match percentage
        - Key alignment points with job requirements
        - Competitive advantages highlighted

        **5. Recommendations**
        - Additional skills to develop for this role
        - Interview preparation focus areas
        - Portfolio projects to emphasize

        Keep the summary concise but comprehensive.
        """
        
        changes_response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": "You are a resume optimization analyst. Provide clear, actionable summaries of improvements made."
                },
                {"role": "user", "content": changes_prompt}
            ],
            max_tokens=700,
            temperature=0.5
        )
        
        optimization_summary = changes_response.choices[0].message.content

        # Increment user's optimization count if authenticated
        if current_user_id:
            user = db.query(User).filter(User.id == current_user_id).first()
            if user:
                user.optimizations_count = (user.optimizations_count or 0) + 1
                db.commit()
        
        return {
            "status": "success",
            "optimized_resume": optimized_resume,
            "optimization_summary": optimization_summary,
            "original_word_count": len(raw_text.split()),
            "optimized_word_count": len(optimized_resume.split()),
            "company_name": request.company_name,
            "position_title": request.position_title,
            "optimization_level": request.optimization_level
        }
        
    except Exception as e:
        return {
            "status": "error",
            "message": f"Error optimizing resume: {str(e)}"
        }
    
@app.post("/api/generate-cover-letter-docx")
async def generate_cover_letter_docx(request: DocxGenerationRequest):
    """Generate a professionally formatted DOCX cover letter"""
    try:
        # Create a new Document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Split content into lines
        lines = request.content.split('\n')
        
        # Track if we're in the header (contact info)
        in_header = True
        
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            
            # Check if this is a header line (contact info at the top)
            if in_header and any(indicator in line.lower() for indicator in ['@', 'linkedin', 'github', 'phone', '|']):
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].bold = True
            # Check if this is the date line
            elif 'dear' in line.lower() or line.startswith('Dear'):
                # Add date first
                date_p = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
                date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                doc.add_paragraph()  # Empty line
                
                # Add the greeting
                p = doc.add_paragraph(line)
                in_header = False
            # Regular paragraph content
            else:
                p = doc.add_paragraph(line)
                if in_header:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if not any(char.islower() for char in line):  # If all caps (like name)
                        p.runs[0].bold = True
                        p.runs[0].font.size = Pt(14)
                in_header = False
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
        
        # Generate filename
        safe_company = "".join(c for c in request.company_name if c.isalnum() or c in (' ', '-', '_')).strip()
        safe_position = "".join(c for c in request.position_title if c.isalnum() or c in (' ', '-', '_')).strip()
        filename = f"Cover_Letter_{safe_company}_{safe_position}.docx".replace(' ', '_')
        
        return FileResponse(
            path=tmp_path,
            filename=filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return {"status": "error", "message": f"Error generating DOCX: {str(e)}"}

@app.post("/api/generate-resume-docx")
async def generate_resume_docx(request: DocxGenerationRequest):
    """Generate a compact, 1-page professional DOCX resume"""
    try:
        # Create a new Document
        doc = Document()
        
        # Set tight margins for maximum space
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Set default font to save space
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)
        
        # Split content into lines
        lines = request.content.split('\n')
        
        # Process content with minimal spacing
        for line in lines:
            line = line.strip()
            if not line:
                continue  # Skip empty lines to save space
            
            # Remove markdown formatting
            line = line.replace('**', '')
            
            # Detect content types and format compactly
            if line.isupper() and len(line) > 5 and len(line) < 30:
                # Name header
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.bold = True
                run.font.size = Pt(14)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.space_after = Pt(6)  # Minimal spacing
                
            elif any(header in line.upper() for header in ['PROFESSIONAL SUMMARY', 'CORE SKILLS', 'EXPERIENCE', 'PROJECTS', 'EDUCATION']):
                # Section headers
                p = doc.add_paragraph()
                run = p.add_run(line.upper())
                run.font.bold = True
                run.font.size = Pt(11)
                p.space_before = Pt(8)
                p.space_after = Pt(4)
                
            elif line.startswith('-') or line.startswith('•'):
                # Bullet points - use hanging indent to save space
                bullet_text = line.lstrip('- •')
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                p.paragraph_format.first_line_indent = Inches(-0.2)
                run = p.add_run(f"• {bullet_text}")
                run.font.size = Pt(9)
                p.space_after = Pt(2)
                
            elif '|' in line and ('@' in line or 'linkedin' in line.lower()):
                # Contact information
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.size = Pt(9)
                run.font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.space_after = Pt(6)
                
            elif len(line) > 80 and not line.startswith('-'):
                # Long paragraphs - professional summary
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.size = Pt(9)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.space_after = Pt(4)
                
            else:
                # Regular content - compact formatting
                p = doc.add_paragraph()
                
                # Check if it's a job title/company line
                if any(word in line for word in ['|', 'Engineer', 'Developer', 'Manager', 'Analyst', 'Intern']):
                    run = p.add_run(line)
                    run.font.bold = True
                    run.font.size = Pt(10)
                else:
                    run = p.add_run(line)
                    run.font.size = Pt(9)
                
                p.space_after = Pt(2)
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
        
        # Generate filename
        safe_company = "".join(c for c in request.company_name if c.isalnum() or c in (' ', '-', '_')).strip()
        safe_position = "".join(c for c in request.position_title if c.isalnum() or c in (' ', '-', '_')).strip()
        
        filename = f"Resume_{safe_company}_{safe_position}_1Page.docx".replace(' ', '_')
        
        return FileResponse(
            path=tmp_path,
            filename=filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return {"status": "error", "message": f"Error generating compact DOCX: {str(e)}"}
    
@app.on_event("startup")
async def startup_event():
    import tempfile
    import glob
    temp_dir = tempfile.gettempdir()
    for file in glob.glob(os.path.join(temp_dir, "tmp*.docx")):
        try:
            os.remove(file)
        except:
            pass

@app.post("/api/auth/register")
async def register_user(user_data: UserCreate, db: Session = Depends(get_db)):
    """Register a new user"""
    try:
        # Check if user already exists
        existing_user = db.query(User).filter(
            (User.email == user_data.email) | (User.username == user_data.username)
        ).first()
        
        if existing_user:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="User with this email or username already exists"
            )
        
        # Create new user
        hashed_password = get_password_hash(user_data.password)
        db_user = User(
            username=user_data.username,
            email=user_data.email,
            hashed_password=hashed_password,
            full_name=user_data.full_name,
            role=user_data.role
        )
        
        db.add(db_user)
        db.commit()
        db.refresh(db_user)
        
        # Create user profile
        user_profile = UserProfile(user_id=db_user.id)
        db.add(user_profile)
        db.commit()
        
        # Create access token
        access_token = create_access_token(data={"sub": str(db_user.id)})
        
        return {
            "access_token": access_token,
            "token_type": "bearer",
            "user": {
                "id": db_user.id,
                "username": db_user.username,
                "email": db_user.email,
                "full_name": db_user.full_name,
                "role": db_user.role
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error creating user: {str(e)}"
        )

@app.post("/api/auth/login")
async def login_user(login_data: UserLogin, db: Session = Depends(get_db)):
    """Login user and return access token"""
    try:
        # Find user by username or email
        user = db.query(User).filter(
            (User.username == login_data.username) | (User.email == login_data.username)
        ).first()
        
        if not user or not verify_password(login_data.password, user.hashed_password):
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Incorrect username/email or password",
                headers={"WWW-Authenticate": "Bearer"},
            )
        
        if not user.is_active:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Inactive user account"
            )
        
        # Create access token
        access_token = create_access_token(
            data={"user_id": user.id}
        )
        
        return {
            "access_token": access_token,
            "token_type": "bearer",
            "user": {
                "id": user.id,
                "username": user.username,
                "email": user.email,
                "full_name": user.full_name,
                "role": user.role,
                "resume_analyses_count": user.resume_analyses_count,
                "cover_letters_count": user.cover_letters_count,
                "optimizations_count": user.optimizations_count,
                "chat_messages_count": user.chat_messages_count
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error during login: {str(e)}"
        )
    
@app.get("/api/auth/me")
async def get_current_user(current_user_id: str = Depends(get_current_user_id), db: Session = Depends(get_db)):
    """Get current user information"""
    try:
        user = db.query(User).filter(User.id == int(current_user_id)).first()
        if not user:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="User not found"
            )
        
        return {
            "id": user.id,
            "username": user.username,
            "email": user.email,
            "full_name": user.full_name,
            "role": user.role,
            "is_active": user.is_active,
            "created_at": user.created_at,
            "resume_analyses_count": user.resume_analyses_count,
            "cover_letters_count": user.cover_letters_count,
            "optimizations_count": user.optimizations_count,
            "chat_messages_count": user.chat_messages_count
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error getting user: {str(e)}"
        )

@app.get("/api/profile")
async def get_user_profile(
    current_user_id: str = Depends(get_current_user_id), 
    db: Session = Depends(get_db)):
    """Get user profile information"""
    try:
        profile = db.query(UserProfile).filter(UserProfile.user_id == int(current_user_id)).first()
        if not profile:
            # Create empty profile if it doesn't exist
            profile = UserProfile(user_id=int(current_user_id))
            db.add(profile)
            db.commit()
            db.refresh(profile)
        
        return {
            "current_title": profile.current_title,
            "experience_level": profile.experience_level,
            "industry": profile.industry,
            "location": profile.location,
            "salary_range": profile.salary_range,
            "technical_skills": json.loads(profile.technical_skills) if profile.technical_skills else [],
            "soft_skills": json.loads(profile.soft_skills) if profile.soft_skills else [],
            "certifications": json.loads(profile.certifications) if profile.certifications else [],
            "languages": json.loads(profile.languages) if profile.languages else [],
            "preferred_roles": json.loads(profile.preferred_roles) if profile.preferred_roles else [],
            "preferred_companies": json.loads(profile.preferred_companies) if profile.preferred_companies else [],
            "remote_preference": profile.remote_preference,
            "willing_to_relocate": profile.willing_to_relocate,
            "linkedin_url": profile.linkedin_url,
            "github_url": profile.github_url,
            "portfolio_url": profile.portfolio_url,
            "updated_at": profile.updated_at
        }
        
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error getting profile: {str(e)}"
        )

@app.put("/api/profile")
async def update_user_profile(
    profile_data: dict,
    current_user_id: str = Depends(get_current_user_id), 
    db: Session = Depends(get_db)
):
    """Update user profile information"""
    try:
        user_id = int(current_user_id)
        profile = db.query(UserProfile).filter(UserProfile.user_id == user_id).first()
        
        if not profile:
            profile = UserProfile(user_id=user_id)
            db.add(profile)
        
        # Update only fields that exist in the model
        for field, value in profile_data.items():
            if hasattr(profile, field):
                setattr(profile, field, value)
        
        db.commit()
        db.refresh(profile)
        
        return {"message": "Profile updated successfully"}
        
    except Exception as e:
        print(f"Error updating profile: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error updating profile: {str(e)}"
        )
    
@app.patch("/api/profile")
async def patch_user_profile(
    profile_data: ProfileUpdate,
    current_user_id: str = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    """Partially update user profile (PATCH method)"""
    return await update_user_profile(profile_data, current_user_id, db)

@app.patch("/api/auth/profile")
async def patch_auth_user_profile(
    profile_data: ProfileUpdate,
    current_user_id: str = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    """Partially update user profile (alternative endpoint)"""
    print(f"PATCH /api/auth/profile called for user: {current_user_id}")
    return await update_user_profile(profile_data, current_user_id, db)
    
@app.post("/api/jobs/fetch")
async def fetch_jobs_from_apis(
    query: str,
    location: str = "United Kingdom",
    max_jobs: int = 50,
    current_user_id: int = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    """Fetch jobs from external APIs and store in database"""
    
    # Fetch jobs from all sources
    jobs_data = job_api_service.fetch_and_store_jobs(query, location, max_jobs)
    
    added_count = 0
    updated_count = 0
    
    for job_data in jobs_data:
        # Check if job already exists
        existing = db.query(JobPosting).filter(
            JobPosting.external_id == job_data['external_id']
        ).first()
        
        if existing:
            # Update existing job
            for key, value in job_data.items():
                setattr(existing, key, value)
            existing.updated_at = datetime.utcnow()
            updated_count += 1
        else:
            # Create new job posting
            # Extract skills from description
            skills = job_api_service.extract_skills_from_description(
                job_data.get('description', '')
            )
            
            job = JobPosting(
                title=job_data['title'],
                company_name=job_data['company_name'],
                company_logo_url=job_data.get('company_logo_url'),
                location=job_data['location'],
                remote_type=job_data['remote_type'],
                description=job_data['description'],
                requirements=job_data.get('requirements', ''),
                salary_min=job_data.get('salary_min'),
                salary_max=job_data.get('salary_max'),
                salary_currency=job_data.get('salary_currency', 'GBP'),
                experience_level=job_data.get('experience_level', 'Mid'),
                employment_type=job_data.get('employment_type', 'full-time'),
                required_skills=skills,
                external_id=job_data['external_id'],
                source=job_data['source'],
                apply_url=job_data.get('apply_url'),
                posted_date=job_data.get('posted_date') or datetime.utcnow(),
                is_active=True
            )
            db.add(job)
            added_count += 1
    
    db.commit()
    
    return {
        "status": "success",
        "jobs_added": added_count,
        "jobs_updated": updated_count,
        "total_fetched": len(jobs_data)
    }


@app.post("/api/job-preferences")
async def create_or_update_job_preferences(
    preferences: dict,
    current_user_id: int = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    """Create or update user job preferences"""
    user_id = int(current_user_id)
    
    existing = db.query(UserJobPreferences).filter(
        UserJobPreferences.user_id == current_user_id
    ).first()
    
    if existing:
        # Update existing preferences
        for key, value in preferences.items():
            if hasattr(existing, key):
                setattr(existing, key, value)
        existing.updated_at = datetime.utcnow()
        db.commit()
        db.refresh(existing)
        return {"status": "updated", "preferences": existing}
    else:
        # Create new preferences
        new_preferences = UserJobPreferences( user_id=user_id)
        for key, value in preferences.items():  
            if hasattr(new_preferences, key):
                setattr(new_preferences, key, value)
        db.add(new_preferences)
        db.commit()
        db.refresh(new_preferences)
        return {"status": "created","preferences": new_preferences}
    
@app.get("/api/job-preferences")
async def get_job_preferences(
    current_user_id: int = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    """Get user job preferences"""
    
    preferences = db.query(UserJobPreferences).filter(
        UserJobPreferences.user_id == current_user_id
    ).first()
    
    if not preferences:
        # Return default preferences
        return {
            "desired_roles": [],
            "remote_preference": "flexible",
            "willing_to_relocate": False
        }
    
    return preferences

@app.get("/api/jobs/recommendations")
async def get_job_recommendations(
    limit: int = 20,
    min_score: float = 0.0,
    current_user_id: int = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    """Get personalized job recommendations with match scores"""
    
    # Find matching jobs using the matching engine
    matches = matching_engine.find_matching_jobs(
        current_user_id, 
        db, 
        limit=limit,
        min_score=min_score
    )
    
    recommendations = []
    
    for job, scores in matches:
        # Check if already saved
        is_saved = db.query(SavedJob).filter(
            and_(
                SavedJob.user_id == current_user_id,
                SavedJob.job_id == job.id
            )
        ).first() is not None
        
        # Check if already applied
        has_applied = db.query(JobApplication).filter(
            and_(
                JobApplication.user_id == current_user_id,
                JobApplication.job_id == job.id
            )
        ).first() is not None
        
        recommendations.append({
            "job": {
                "id": job.id,
                "title": job.title,
                "company_name": job.company_name,
                "company_logo_url": job.company_logo_url,
                "location": job.location,
                "remote_type": job.remote_type,
                "salary_min": job.salary_min,
                "salary_max": job.salary_max,
                "experience_level": job.experience_level,
                "employment_type": job.employment_type,
                "description": job.description,
                "required_skills": job.required_skills,
                "company_size": job.company_size,
                "industry": job.industry,
                "posted_date": job.posted_date.isoformat() if job.posted_date else None,
                "is_saved": is_saved,
                "apply_url": job.apply_url,
                "has_applied": has_applied
            },
            "match_score": round(scores['overall_score'], 1),
            "scores": {
                "skills": round(scores['skills_score'], 1),
                "experience": round(scores['experience_score'], 1),
                "location": round(scores['location_score'], 1),
                "salary": round(scores['salary_score'], 1),
                "company": round(scores['company_score'], 1)
            }
        })
    
    return {"recommendations": recommendations}

@app.post("/api/jobs/{job_id}/save")
async def save_job(
    job_id: int,
    current_user_id: int = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    """Save a job for later"""
    
    # Verify job exists
    job = db.query(JobPosting).filter(JobPosting.id == job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    # Check if already saved
    existing = db.query(SavedJob).filter(
        and_(
            SavedJob.user_id == current_user_id,
            SavedJob.job_id == job_id
        )
    ).first()
    
    if existing:
        return {"message": "Job already saved", "saved": True}
    
    # Create saved job
    saved_job = SavedJob(
        user_id=current_user_id,
        job_id=job_id
    )
    db.add(saved_job)
    db.commit()
    
    return {"message": "Job saved successfully", "saved": True}

@app.delete("/api/jobs/{job_id}/unsave")
async def unsave_job(
    job_id: int,
    current_user_id: int = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    """Remove a saved job"""
    
    saved_job = db.query(SavedJob).filter(
        and_(
            SavedJob.user_id == current_user_id,
            SavedJob.job_id == job_id
        )
    ).first()
    
    if not saved_job:
        raise HTTPException(status_code=404, detail="Saved job not found")
    
    db.delete(saved_job)
    db.commit()
    
    return {"message": "Job removed from saved", "saved": False}

@app.get("/api/jobs/saved")
async def get_saved_jobs(
    current_user_id: int = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    """Get all saved jobs"""
    
    saved_jobs = db.query(SavedJob).filter(
        SavedJob.user_id == current_user_id
    ).order_by(desc(SavedJob.saved_at)).all()
    
    result = []
    for saved in saved_jobs:
        job = db.query(JobPosting).filter(JobPosting.id == saved.job_id).first()
        if job and job.is_active:
            result.append({
                "id": job.id,
                "title": job.title,
                "company_name": job.company_name,
                "location": job.location,
                "remote_type": job.remote_type,
                "salary_min": job.salary_min,
                "salary_max": job.salary_max,
                "saved_at": saved.saved_at.isoformat()
            })
    
    return {"saved_jobs": result}

@app.post("/api/applications")
async def create_application(
    application: JobApplicationCreate,
    current_user_id: int = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    """Submit a job application"""
    
    # Verify job exists
    job = db.query(JobPosting).filter(JobPosting.id == application.job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    # Check if already applied
    existing = db.query(JobApplication).filter(
        and_(
            JobApplication.user_id == current_user_id,
            JobApplication.job_id == application.job_id
        )
    ).first()
    
    if existing:
        raise HTTPException(
            status_code=400,
            detail="You have already applied to this job"
        )
    
    # Create application
    new_application = JobApplication(
        user_id=current_user_id,
        job_id=application.job_id,
        cover_letter_used=application.cover_letter,
        resume_version_used=application.resume_version,
        status="applied",
        applied_date=datetime.utcnow()
    )

    db.add(new_application)
    db.commit()
    db.refresh(new_application)
    
    return {
        "message": "Application submitted successfully",
        "application_id": new_application.id,
        "status": "applied"
    }


@app.get("/api/applications")
async def get_applications(
    status: Optional[str] = None,
    current_user_id: str= Depends(get_current_user_id),  # ✅ Using str instead of Optional
    db: Session = Depends(get_db)
):
    """Get all user applications"""
    # ✅ Proper check for None
    user_id = int(current_user_id)  # ✅ Convert to int
    
    query = db.query(JobApplication).filter(
        JobApplication.user_id == user_id  # ✅ Use user_id here
    )
    
    if status:
        query = query.filter(JobApplication.status == status)
    
    applications = query.order_by(desc(JobApplication.applied_date)).all()
    
    result = []
    for app in applications:
        job = db.query(JobPosting).filter(JobPosting.id == app.job_id).first()
        if job:
            next_step = "Awaiting response"
            if hasattr(app, 'user_notes') and app.user_notes:
                next_step = app.user_notes
            elif hasattr(app, 'interview_scheduled') and app.interview_scheduled:
                next_step = f"Interview scheduled for {app.interview_scheduled.strftime('%Y-%m-%d')}"
            
            result.append({
                "id": app.id,
                "jobTitle": job.title,
                "company": job.company_name,
                "status": app.status,
                "appliedDate": app.applied_date.strftime("%Y-%m-%d"),
                "nextStep": next_step,
                "interview_date": app.interview_scheduled.isoformat() if hasattr(app, 'interview_scheduled') and app.interview_scheduled else None
            })
    
    return result

@app.patch("/api/applications/{application_id}")
async def update_application(
    application_id: int,
    update: JobApplicationUpdate,
    current_user_id: int = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    """Update application status and notes"""
    
    application = db.query(JobApplication).filter(
        and_(
            JobApplication.id == application_id,
            JobApplication.user_id == current_user_id
        )
    ).first()
    
    if not application:
        raise HTTPException(status_code=404, detail="Application not found")
    
    # Update fields
    if update.status:
        application.status = update.status
    if update.notes:
        application.notes = update.notes
    if update.interview_date:
        application.interview_date = datetime.fromisoformat(update.interview_date)
    
    application.updated_at = datetime.utcnow()
    db.commit()
    
    return {
        "message": "Application updated successfully",
        "status": application.status
    }

@app.get("/api/jobs/search")
async def search_jobs(
    query: Optional[str] = None,
    location: Optional[str] = None,
    remote_type: Optional[str] = None,
    min_salary: Optional[int] = None,
    experience_level: Optional[str] = None,
    skip: int = 0,
    limit: int = 20,
    db: Session = Depends(get_db)
):
    """Search jobs with filters"""
    
    job_query = db.query(JobPosting).filter(JobPosting.is_active == True)
    
    if query:
        job_query = job_query.filter(
            or_(
                JobPosting.title.ilike(f"%{query}%"),
                JobPosting.company_name.ilike(f"%{query}%"),
                JobPosting.description.ilike(f"%{query}%")
            )
        )
    
    if location:
        job_query = job_query.filter(
            JobPosting.location.ilike(f"%{location}%")
        )
    
    if remote_type:
        job_query = job_query.filter(JobPosting.remote_type == remote_type)
    
    if min_salary:
        job_query = job_query.filter(JobPosting.salary_max >= min_salary)
    
    if experience_level:
        job_query = job_query.filter(
            JobPosting.experience_level == experience_level
        )
    
    total = job_query.count()
    jobs = job_query.order_by(desc(JobPosting.posted_date)).offset(skip).limit(limit).all()
    
    return {
        "jobs": jobs,
        "total": total,
        "page": skip // limit + 1,
        "pages": (total + limit - 1) // limit
    }

@app.get("/api/jobs/{job_id}")
async def get_job_details(
    job_id: int,
    db: Session = Depends(get_db)
):
    """Get detailed job information"""
    
    job = db.query(JobPosting).filter(JobPosting.id == job_id).first()
    
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    return job

@app.get("/api/resumes")
async def get_resumes(
    current_user_id: str = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    """Get all user resumes"""
    user_id = int(current_user_id)
    
    resumes = db.query(Resume).filter(Resume.user_id == user_id).all()
    
    return [
        {
            "id": resume.id,
            "title": resume.title,
            "fileName": resume.file_name,
            "uploadDate": resume.upload_date.isoformat(),
            "isDefault": resume.is_default
        }
        for resume in resumes
    ]

@app.post("/api/resumes/upload")
async def upload_resume(
    file: UploadFile = File(...),
    title: str = Form(...),
    current_user_id: str = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    """Upload a new resume"""
    user_id = int(current_user_id)
    
    # Save file
    file_name = f"{user_id}_{int(datetime.utcnow().timestamp())}_{file.filename}"
    file_path = os.path.join(UPLOAD_DIR, file_name)
    
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    # Create database record
    resume = Resume(
        user_id=user_id,
        title=title,
        file_name=file.filename,
        file_path=file_path,
        is_default=False
    )
    db.add(resume)
    db.commit()
    
    return {"message": "Resume uploaded successfully", "id": resume.id}

@app.patch("/api/resumes/{resume_id}/set-default")
async def set_default_resume(
    resume_id: int,
    current_user_id: str = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    """Set a resume as default"""
    user_id = int(current_user_id)
    
    # Remove default from all user resumes
    db.query(Resume).filter(Resume.user_id == user_id).update({"is_default": False})
    
    # Set new default
    resume = db.query(Resume).filter(
        Resume.id == resume_id,
        Resume.user_id == user_id
    ).first()
    
    if resume:
        resume.is_default = True
        db.commit()
    
    return {"message": "Default resume updated"}

@app.delete("/api/resumes/{resume_id}")
async def delete_resume(
    resume_id: int,
    current_user_id: str = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    """Delete a resume"""
    user_id = int(current_user_id)
    
    resume = db.query(Resume).filter(
        Resume.id == resume_id,
        Resume.user_id == user_id
    ).first()
    
    if resume:
        # Delete file
        if os.path.exists(resume.file_path):
            os.remove(resume.file_path)
        
        # Delete database record
        db.delete(resume)
        db.commit()
    
    return {"message": "Resume deleted"}

@app.get("/api/resumes/{resume_id}/download")
async def download_resume(
    resume_id: int,
    current_user_id: str = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    """Download a resume"""
    from fastapi.responses import FileResponse
    
    user_id = int(current_user_id)
    
    resume = db.query(Resume).filter(
        Resume.id == resume_id,
        Resume.user_id == user_id
    ).first()
    
    if not resume or not os.path.exists(resume.file_path):
        raise HTTPException(status_code=404, detail="Resume not found")
    
    return FileResponse(resume.file_path, filename=resume.file_name)


#Debugging below this line    

@app.get("/api/admin/test-adzuna")
async def test_adzuna_direct():
    """Test Adzuna API with direct call"""
    import requests
    
    app_id = os.getenv("ADZUNA_APP_ID")
    app_key = os.getenv("ADZUNA_API_KEY")
    
    # Test with simple search
    url = "https://api.adzuna.com/v1/api/jobs/gb/search/1"
    
    params = {
        "app_id": app_id,
        "app_key": app_key,
        "results_per_page": 5,
        "what": "developer"
    }
    
    try:
        print(f"Testing URL: {url}")
        print(f"With params: {params}")
        
        response = requests.get(url, params=params, timeout=30)
        
        return {
            "status_code": response.status_code,
            "url": response.url,
            "response_preview": response.text[:500],
            "success": response.status_code == 200,
            "jobs_found": len(response.json().get("results", [])) if response.status_code == 200 else 0
        }
    except Exception as e:
        return {
            "error": str(e),
            "traceback": traceback.format_exc()
        }

@app.get("/api/admin/check-jobs")
async def check_jobs_in_database(db: Session = Depends(get_db)):
    """Check how many jobs are in the database"""
    total_jobs = db.query(JobPosting).count()
    active_jobs = db.query(JobPosting).filter(JobPosting.is_active == True).count()
    mock_jobs = db.query(JobPosting).filter(JobPosting.source == "mock").count()
    adzuna_jobs = db.query(JobPosting).filter(JobPosting.source == "adzuna").count()
    
    # Get sample jobs
    sample_jobs = db.query(JobPosting).limit(5).all()
    
    return {
        "total_jobs": total_jobs,
        "active_jobs": active_jobs,
        "mock_jobs": mock_jobs,
        "adzuna_jobs": adzuna_jobs,
        "sample_jobs": [
            {
                "id": job.id,
                "title": job.title,
                "company": job.company_name,
                "source": job.source,
                "skills": job.required_skills
            }
            for job in sample_jobs
        ]
    }

@app.get("/api/admin/debug-matching/{user_id}")
async def debug_matching(user_id: int, db: Session = Depends(get_db)):
    """Debug matching for a specific user"""
    user = db.query(User).filter(User.id == user_id).first()
    user_profile = db.query(UserProfile).filter(UserProfile.user_id == user_id).first()
    user_preferences = db.query(UserJobPreferences).filter(UserJobPreferences.user_id == user_id).first()
    
    # Get all jobs
    all_jobs = db.query(JobPosting).filter(JobPosting.is_active == True).all()
    
    # Try to get matches
    try:
        matches = matching_engine.find_matching_jobs(user_id, db, limit=20, min_score=0)
        match_count = len(matches)
        sample_match = matches[0] if matches else None
    except Exception as e:
        match_count = 0
        sample_match = str(e)
    
    return {
        "user_id": user_id,
        "user_exists": user is not None,
        "profile_exists": user_profile is not None,
        "preferences_exists": user_preferences is not None,
        "total_active_jobs": len(all_jobs),
        "matches_found": match_count,
        "sample_match_score": sample_match[1]['overall_score'] if sample_match and isinstance(sample_match, tuple) else None,
        "error": sample_match if isinstance(sample_match, str) else None,
        "profile_data": {
            "skills": user_profile.technical_skills if user_profile else None,
            "experience": user_profile.experience_level if user_profile else None
        },
        "preferences_data": {
            "remote": user_preferences.remote_preference if user_preferences else None,
            "salary": user_preferences.minimum_salary if user_preferences else None
        } if user_preferences else None
    }


# Email Alerts  feature

@app.post("/api/job-alerts")
async def create_job_alert(
    email: str,
    min_match_score: int = 80,
    current_user_id: str = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    """Create job alert"""
    user_id = int(current_user_id)
    
    alert = JobAlert(
        user_id=user_id,
        email=email,
        min_match_score=min_match_score,
        is_active=True
    )
    db.add(alert)
    db.commit()
    
    return {"message": "Job alert created successfully"}
    

@app.get("/api/job-alerts")
async def get_job_alerts(
    current_user_id: str = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    """Get user's job alerts"""
    user_id = int(current_user_id)
    alerts = db.query(JobAlert).filter(JobAlert.user_id == user_id).all()
    return alerts

@app.delete("/api/job-alerts/{alert_id}")
async def delete_job_alert(
    alert_id: int,
    current_user_id: str = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    """Delete job alert"""
    user_id = int(current_user_id)
    alert = db.query(JobAlert).filter(
        JobAlert.id == alert_id,
        JobAlert.user_id == user_id
    ).first()
    
    if alert:
        db.delete(alert)
        db.commit()
    
    return {"message": "Alert deleted"}


# Background job scheduler
def fetch_jobs_and_send_alerts():
    """Run daily to fetch jobs and send alerts"""
    from database import SessionLocal
    db = SessionLocal()
    
    try:
        # Fetch new jobs
        print("Fetching new jobs...")
        queries = ["software developer", "frontend developer", "backend developer"]
        
        for query in queries:
            jobs_data = job_api_service.fetch_and_store_jobs(query, "United Kingdom", 20)
            
            for job_data in jobs_data:
                existing = db.query(JobPosting).filter(
                    JobPosting.external_id == job_data['external_id']
                ).first()
                
                if not existing:
                    skills = job_api_service.extract_skills_from_description(
                        job_data.get('description', '')
                    )
                    
                    job = JobPosting(
                        title=job_data['title'],
                        company_name=job_data['company_name'],
                        location=job_data['location'],
                        remote_type=job_data['remote_type'],
                        description=job_data['description'],
                        salary_min=job_data.get('salary_min'),
                        salary_max=job_data.get('salary_max'),
                        salary_currency='GBP',
                        experience_level='Mid',
                        employment_type=job_data['employment_type'],
                        required_skills=skills,
                        external_id=job_data['external_id'],
                        source=job_data['source'],
                        apply_url=job_data.get('apply_url'),
                        posted_date=job_data.get('posted_date') or datetime.utcnow(),
                        is_active=True
                    )
                    db.add(job)
            
            db.commit()
        
        # Send alerts to users
        print("Sending job alerts...")
        alerts = db.query(JobAlert).filter(JobAlert.is_active == True).all()
        
        for alert in alerts:
            matches = matching_engine.find_matching_jobs(
                alert.user_id, db, limit=10, min_score=alert.min_match_score
            )
            
            if matches:
                jobs_to_send = [
                    {
                        'title': job.title,
                        'company': job.company_name,
                        'location': job.location,
                        'match_score': round(scores['overall_score']),
                        'apply_url': job.apply_url or '#'
                    }
                    for job, scores in matches
                ]
                
                import asyncio
                asyncio.run(send_job_alert_email(alert.email, jobs_to_send))
                
                alert.last_sent = datetime.utcnow()
                db.commit()
        
        print("Job fetch and alerts completed")
    
    finally:
        db.close()

# Schedule daily job fetching at 9 AM
scheduler = BackgroundScheduler()
scheduler.add_job(fetch_jobs_and_send_alerts, 'cron', hour=21, minute=30)
scheduler.start()

# Shutdown scheduler on exit
atexit.register(lambda: scheduler.shutdown())




#Debugging codes below this line 
@app.get("/api/debug/all-jobs")
async def debug_all_jobs(
    current_user_id: str = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    """Debug: Show all jobs with their match scores"""
    user_id = int(current_user_id)
    
    # Get all active jobs
    jobs = db.query(JobPosting).filter(JobPosting.is_active == True).order_by(desc(JobPosting.posted_date)).limit(20).all()
    
    result = []
    for job in jobs:
        # Calculate match score
        try:
            user = db.query(User).filter(User.id == user_id).first()
            user_profile = db.query(UserProfile).filter(UserProfile.user_id == user_id).first()
            user_preferences = db.query(UserJobPreferences).filter(UserJobPreferences.user_id == user_id).first()
            
            overall_score, detailed_scores = matching_engine.calculate_match_score(
                user, job, user_profile, user_preferences
            )
            
            result.append({
                "id": job.id,
                "title": job.title,
                "company": job.company_name,
                "posted_date": job.posted_date.isoformat() if job.posted_date else None,
                "match_score": round(overall_score, 1),
                "scores": detailed_scores
            })
        except Exception as e:
            result.append({
                "id": job.id,
                "title": job.title,
                "error": str(e)
            })
    
    return result

@app.get("/api/admin/debug-jobs")
async def debug_jobs_simple(db: Session = Depends(get_db)):
    """Debug: Show recent jobs"""
    jobs = db.query(JobPosting).filter(JobPosting.is_active == True).order_by(desc(JobPosting.posted_date)).limit(10).all()
    
    return [
        {
            "id": job.id,
            "title": job.title,
            "company": job.company_name,
            "posted_date": str(job.posted_date),
            "skills": job.required_skills
        }
        for job in jobs
    ]

# @app.get("/api/debug/check-auth")
# async def check_auth(authorization: Optional[str] = Header(None)):
#     """Debug authentication"""
#     return {
#         "authorization_header": authorization,
#         "has_bearer": "Bearer" in (authorization or ""),
#         "token_preview": authorization[:50] if authorization else None
#     }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
